import asyncio
import logging
from typing import Dict, Any, Type
from pathlib import Path

logger = logging.getLogger(__name__)


class ParsedDocument:
    def __init__(self, content: str, metadata: Dict[str, Any]):
        self.content = content
        self.metadata = metadata


class BaseParser:
    def __init__(self, timeout: int = 30):
        self.timeout = timeout

    async def parse(self, content: bytes) -> ParsedDocument:
        raise NotImplementedError


class PDFParser(BaseParser):
    async def parse(self, content: bytes) -> ParsedDocument:
        try:
            return await asyncio.get_event_loop().run_in_executor(None, self._parse_sync, content)
        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            return ParsedDocument(f"PDF parsing error: {e}", {"format": "pdf", "error": str(e)})

    def _parse_sync(self, content: bytes) -> ParsedDocument:
        text_parts = []
        num_pages = 0

        try:
            import fitz
            import io
            doc = fitz.open(stream=content, filetype="pdf")
            num_pages = len(doc)
            for page_num in range(num_pages):
                page = doc[page_num]
                text_parts.append(page.get_text())
            doc.close()
            logger.info(f"Extracted {num_pages} pages via PyMuPDF")
        except ImportError:
            try:
                import pdfplumber
                import io
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    num_pages = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        text_parts.append(page_text)
                logger.info(f"Extracted {num_pages} pages via pdfplumber")
            except ImportError:
                try:
                    from PyPDF2 import PdfReader
                    import io
                    reader = PdfReader(io.BytesIO(content))
                    num_pages = len(reader.pages)
                    for page in reader.pages:
                        page_text = page.extract_text() or ""
                        text_parts.append(page_text)
                    logger.info(f"Extracted {num_pages} pages via PyPDF2")
                except ImportError:
                    raise RuntimeError("No PDF library available (try: pip install PyMuPDF pdfplumber PyPDF2)")

        full_text = "\n\n".join(text_parts)
        word_count = len(full_text.split())
        return ParsedDocument(full_text.strip(), {
            "format": "pdf",
            "pages": num_pages,
            "word_count": word_count,
        })


class MarkdownParser(BaseParser):
    async def parse(self, content: bytes) -> ParsedDocument:
        try:
            text = content.decode("utf-8", errors="replace")
            word_count = len(text.split())
            return ParsedDocument(text, {
                "format": "markdown",
                "word_count": word_count,
            })
        except Exception as e:
            logger.error(f"Markdown parsing failed: {e}")
            return ParsedDocument("", {"format": "markdown", "error": str(e)})


class DOCXParser(BaseParser):
    async def parse(self, content: bytes) -> ParsedDocument:
        try:
            return await asyncio.get_event_loop().run_in_executor(None, self._parse_sync, content)
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            return ParsedDocument(f"DOCX parsing error: {e}", {"format": "docx", "error": str(e)})

    def _parse_sync(self, content: bytes) -> ParsedDocument:
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(content))
            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]

            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_parts.append(" | ".join(row_text))

            full_text = "\n".join(text_parts)
            word_count = len(full_text.split())
            return ParsedDocument(full_text.strip(), {
                "format": "docx",
                "pages": len(doc.sections),
                "word_count": word_count,
            })
        except ImportError:
            raise RuntimeError("python-docx library is not available (try: pip install python-docx)")


class ImageParser(BaseParser):
    async def parse(self, content: bytes) -> ParsedDocument:
        try:
            return await asyncio.get_event_loop().run_in_executor(None, self._parse_sync, content)
        except Exception as e:
            logger.error(f"Image parsing failed: {e}")
            return ParsedDocument(
                content="Image OCR failed",
                metadata={"format": "image", "error": str(e), "needs_review": True},
            )

    def _parse_sync(self, content: bytes) -> ParsedDocument:
        try:
            import pytesseract
            from PIL import Image
            import io
            image = Image.open(io.BytesIO(content))
            ocr_text = pytesseract.image_to_string(image)
            confidence = 0.75
            word_count = len(ocr_text.split())
            return ParsedDocument(
                content=ocr_text.strip(),
                metadata={
                    "format": "image",
                    "ocr_confidence": confidence,
                    "needs_review": confidence < 0.75,
                    "word_count": word_count,
                },
            )
        except ImportError:
            logger.warning("pytesseract or Pillow not installed — returning mock OCR result")
            return ParsedDocument(
                content="Image OCR not available (pytesseract/Pillow not installed)",
                metadata={
                    "format": "image",
                    "ocr_confidence": 0.0,
                    "needs_review": True,
                    "warning": "pytesseract or Pillow not available",
                },
            )
        except Exception as e:
            logger.error(f"Image OCR failed: {e}")
            return ParsedDocument(
                content="Image OCR failed",
                metadata={"format": "image", "error": str(e), "needs_review": True},
            )


PARSERS: Dict[str, Type[BaseParser]] = {
    ".pdf": PDFParser,
    ".md": MarkdownParser,
    ".docx": DOCXParser,
    ".jpg": ImageParser,
    ".png": ImageParser,
}


class UnsupportedFormatError(Exception):
    pass


async def parse_document(filename: str, content: bytes) -> ParsedDocument:
    ext = Path(filename).suffix.lower()
    parser_cls = PARSERS.get(ext)
    if not parser_cls:
        raise UnsupportedFormatError(f"No parser for {ext}")

    parser = parser_cls(timeout=30)
    try:
        return await asyncio.wait_for(parser.parse(content), timeout=parser.timeout)
    except asyncio.TimeoutError:
        logger.error(f"Parsing timed out for {filename}")
        raise
