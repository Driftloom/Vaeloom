import io
import sys
import hashlib
import asyncio
from unittest.mock import MagicMock, AsyncMock

import pytest


@pytest.fixture(autouse=True)
def _clean_parser_modules(monkeypatch):
    # xdist reuses workers across files; ensure parser sys.modules is clean
    # so previous file's monkeypatch (e.g. fitz=None) doesn't leak. Each test
    # then sets exactly what it needs via monkeypatch.setitem.
    for mod in ("fitz", "pdfplumber", "PyPDF2", "docx", "pytesseract", "PIL", "PIL.Image"):
        monkeypatch.delitem(sys.modules, mod, raising=False)
    yield
    for mod in ("fitz", "pdfplumber", "PyPDF2", "docx", "pytesseract", "PIL", "PIL.Image"):
        monkeypatch.delitem(sys.modules, mod, raising=False)


class TestParsers:

    # --- BaseParser ---

    async def test_base_parser_not_implemented(self):
        from api.ingestion.parsers import BaseParser
        parser = BaseParser()
        with pytest.raises(NotImplementedError):
            await parser.parse(b"test")

    # --- PDFParser ---

    async def test_pdf_parser_fitz_success(self, monkeypatch):
        from api.ingestion.parsers import PDFParser

        mock_fitz = MagicMock()
        mock_doc = MagicMock()
        mock_doc.__len__.return_value = 2
        mock_pages = [MagicMock(), MagicMock()]
        mock_pages[0].get_text.return_value = "Page one content"
        mock_pages[1].get_text.return_value = "Page two content"
        mock_doc.__getitem__.side_effect = lambda i: mock_pages[i]
        mock_fitz.open.return_value = mock_doc
        monkeypatch.setitem(sys.modules, "fitz", mock_fitz)

        parser = PDFParser()
        result = await parser.parse(b"fake pdf")
        assert "Page one content" in result.content
        assert "Page two content" in result.content
        assert result.metadata["format"] == "pdf"
        assert result.metadata["pages"] == 2

    async def test_pdf_parser_pdfplumber_success(self, monkeypatch):
        from api.ingestion.parsers import PDFParser

        monkeypatch.setitem(sys.modules, "fitz", None)

        mock_pdfplumber = MagicMock()
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Plumber extracted"
        mock_pdf = MagicMock()
        mock_pdf.pages = [mock_page]
        mock_pdf.__len__.return_value = 1
        mock_pdfplumber.open.return_value.__enter__.return_value = mock_pdf
        monkeypatch.setitem(sys.modules, "pdfplumber", mock_pdfplumber)

        parser = PDFParser()
        result = await parser.parse(b"fake pdf")
        assert "Plumber extracted" in result.content
        assert result.metadata["format"] == "pdf"
        assert result.metadata["pages"] == 1

    async def test_pdf_parser_pypdf2_success(self, monkeypatch):
        from api.ingestion.parsers import PDFParser

        monkeypatch.setitem(sys.modules, "fitz", None)
        monkeypatch.setitem(sys.modules, "pdfplumber", None)

        mock_PyPDF2 = MagicMock()
        mock_reader = MagicMock()
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "PyPDF2 extracted"
        mock_reader.pages = [mock_page]
        mock_PyPDF2.PdfReader.return_value = mock_reader
        monkeypatch.setitem(sys.modules, "PyPDF2", mock_PyPDF2)

        parser = PDFParser()
        result = await parser.parse(b"fake pdf")
        assert "PyPDF2 extracted" in result.content
        assert result.metadata["format"] == "pdf"
        assert result.metadata["pages"] == 1

    async def test_pdf_parser_all_missing(self, monkeypatch):
        from api.ingestion.parsers import PDFParser

        monkeypatch.setitem(sys.modules, "fitz", None)
        monkeypatch.setitem(sys.modules, "pdfplumber", None)
        monkeypatch.setitem(sys.modules, "PyPDF2", None)

        parser = PDFParser()
        result = await parser.parse(b"fake pdf")
        assert "error" in result.metadata
        assert "No PDF library" in result.content

    async def test_pdf_parser_exception_in_fitz(self, monkeypatch):
        from api.ingestion.parsers import PDFParser

        mock_fitz = MagicMock()
        mock_fitz.open.side_effect = ValueError("fitz corrupt")
        monkeypatch.setitem(sys.modules, "fitz", mock_fitz)

        parser = PDFParser()
        result = await parser.parse(b"fake pdf")
        assert "error" in result.metadata
        assert "fitz corrupt" in result.content

    # --- MarkdownParser ---

    async def test_markdown_parser_success(self):
        from api.ingestion.parsers import MarkdownParser

        parser = MarkdownParser()
        result = await parser.parse(b"# Hello\n\nWorld text")
        assert "# Hello" in result.content
        assert result.metadata["format"] == "markdown"
        assert result.metadata["word_count"] == 4

    async def test_markdown_parser_decode_error(self, monkeypatch):
        from api.ingestion.parsers import MarkdownParser

        class BadBytes:
            def decode(self, encoding="utf-8", errors="replace"):
                raise ValueError("decode failed")

        parser = MarkdownParser()
        result = await parser.parse(BadBytes())
        assert result.metadata.get("error") == "decode failed"

    # --- DOCXParser ---

    async def test_docx_parser_success(self):
        from api.ingestion.parsers import DOCXParser
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello World")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        buf = io.BytesIO()
        doc.save(buf)

        parser = DOCXParser()
        result = await parser.parse(buf.getvalue())
        assert "Hello World" in result.content
        assert "A | B" in result.content
        assert result.metadata["format"] == "docx"
        assert result.metadata["pages"] == 1

    async def test_docx_parser_import_error(self, monkeypatch):
        from api.ingestion.parsers import DOCXParser

        monkeypatch.setitem(sys.modules, "docx", None)

        parser = DOCXParser()
        result = await parser.parse(b"fake docx")
        assert "error" in result.metadata
        assert "python-docx" in result.content

    async def test_docx_parser_exception(self, monkeypatch):
        from api.ingestion.parsers import DOCXParser

        mock_docx = MagicMock()
        mock_docx.Document.side_effect = ValueError("docx error")
        monkeypatch.setitem(sys.modules, "docx", mock_docx)

        parser = DOCXParser()
        result = await parser.parse(b"fake docx")
        assert "error" in result.metadata

    # --- ImageParser ---

    async def test_image_parser_success(self, monkeypatch):
        from api.ingestion.parsers import ImageParser

        mock_pytesseract = MagicMock()
        mock_pytesseract.image_to_string.return_value = "OCR recognized text"
        monkeypatch.setitem(sys.modules, "pytesseract", mock_pytesseract)

        mock_PIL = MagicMock()
        mock_PIL.Image = MagicMock()
        monkeypatch.setitem(sys.modules, "PIL", mock_PIL)

        parser = ImageParser()
        result = await parser.parse(b"fake image")
        assert "OCR recognized text" in result.content
        assert result.metadata["format"] == "image"
        assert result.metadata["ocr_confidence"] == 0.75
        assert result.metadata["needs_review"] is False

    async def test_image_parser_import_error(self, monkeypatch):
        from api.ingestion.parsers import ImageParser

        monkeypatch.setitem(sys.modules, "pytesseract", None)
        monkeypatch.setitem(sys.modules, "PIL", None)

        parser = ImageParser()
        result = await parser.parse(b"fake image")
        assert "OCR not available" in result.content
        assert result.metadata["ocr_confidence"] == 0.0
        assert result.metadata["needs_review"] is True

    async def test_image_parser_ocr_exception(self, monkeypatch):
        from api.ingestion.parsers import ImageParser

        mock_pytesseract = MagicMock()
        mock_pytesseract.image_to_string.side_effect = RuntimeError("OCR crashed")
        monkeypatch.setitem(sys.modules, "pytesseract", mock_pytesseract)

        mock_PIL = MagicMock()
        mock_PIL.Image = MagicMock()
        monkeypatch.setitem(sys.modules, "PIL", mock_PIL)

        parser = ImageParser()
        result = await parser.parse(b"fake image")
        assert "OCR failed" in result.content
        assert result.metadata["needs_review"] is True
        assert "OCR crashed" in result.metadata["error"]

    async def test_image_parser_outer_exception(self, monkeypatch):
        from api.ingestion.parsers import ImageParser

        def no_loop():
            raise RuntimeError("no event loop")

        monkeypatch.setattr(asyncio, "get_event_loop", no_loop)

        parser = ImageParser()
        result = await parser.parse(b"fake image")
        assert "error" in result.metadata
        assert "no event loop" in result.metadata["error"]

    # --- parse_document ---

    async def test_parse_document_known_ext(self, monkeypatch):
        from api.ingestion.parsers import parse_document

        mock_fitz = MagicMock()
        mock_doc = MagicMock()
        mock_doc.__len__.return_value = 1
        mock_page = MagicMock()
        mock_page.get_text.return_value = "Parsed via dispatch"
        mock_doc.__getitem__.return_value = mock_page
        mock_fitz.open.return_value = mock_doc
        monkeypatch.setitem(sys.modules, "fitz", mock_fitz)

        result = await parse_document("report.pdf", b"data")
        assert result.metadata["format"] == "pdf"
        assert "Parsed via dispatch" in result.content

    async def test_parse_document_unknown_ext(self):
        from api.ingestion.parsers import parse_document, UnsupportedFormatError

        with pytest.raises(UnsupportedFormatError, match="No parser for"):
            await parse_document("unknown.xyz", b"data")

    async def test_parse_document_timeout(self, monkeypatch):
        import api.ingestion.parsers as parsers_module
        from api.ingestion.parsers import (
            parse_document,
            BaseParser,
            ParsedDocument,
        )

        class TimeoutParser(BaseParser):
            timeout = 0.001

            async def parse(self, content):
                await asyncio.sleep(100)
                return ParsedDocument("", {})

        monkeypatch.setitem(parsers_module.PARSERS, ".pdf", TimeoutParser)

        with pytest.raises(asyncio.TimeoutError):
            await parse_document("test.pdf", b"data")


class TestDedup:

    async def test_compute_content_hash(self):
        from api.ingestion.dedup import compute_content_hash

        result = compute_content_hash(b"hello")
        expected = hashlib.sha256(b"hello").hexdigest()
        assert result == expected

    async def test_check_dedup_version_match(self, monkeypatch):
        from api.ingestion.dedup import check_dedup

        mock_doc = MagicMock()
        mock_doc.id = "doc-uuid-123"
        mock_version = MagicMock()
        mock_version.document_id = "doc-uuid-123"

        mock_version_result = MagicMock()
        mock_version_result.scalar_one_or_none.return_value = mock_version

        mock_doc_result = MagicMock()
        mock_doc_result.scalar_one_or_none.return_value = mock_doc

        mock_session = AsyncMock()
        mock_session.execute.side_effect = [mock_version_result, mock_doc_result]

        mock_cm = AsyncMock()
        mock_cm.__aenter__.return_value = mock_session
        mock_factory = MagicMock(return_value=mock_cm)
        monkeypatch.setattr("api.database.async_session_factory", mock_factory)

        result = await check_dedup("ws1", "hash123", "doc.pdf")
        assert result == "doc-uuid-123"

    async def test_check_dedup_version_orphaned(self, monkeypatch):
        from api.ingestion.dedup import check_dedup

        mock_version = MagicMock()
        mock_version.document_id = "orphaned-doc-id"

        mock_version_result = MagicMock()
        mock_version_result.scalar_one_or_none.return_value = mock_version

        mock_no_doc_result = MagicMock()
        mock_no_doc_result.scalar_one_or_none.return_value = None

        mock_path_result = MagicMock()
        mock_path_result.scalar_one_or_none.return_value = None

        mock_session = AsyncMock()
        mock_session.execute.side_effect = [
            mock_version_result,
            mock_no_doc_result,
            mock_path_result,
        ]

        mock_cm = AsyncMock()
        mock_cm.__aenter__.return_value = mock_session
        mock_factory = MagicMock(return_value=mock_cm)
        monkeypatch.setattr("api.database.async_session_factory", mock_factory)

        result = await check_dedup("ws1", "hash", "orphaned.pdf")
        assert result is None

    async def test_check_dedup_path_match(self, monkeypatch):
        from api.ingestion.dedup import check_dedup

        mock_doc = MagicMock()
        mock_doc.id = "path-match-uuid"

        mock_no_version_result = MagicMock()
        mock_no_version_result.scalar_one_or_none.return_value = None

        mock_path_result = MagicMock()
        mock_path_result.scalar_one_or_none.return_value = mock_doc

        mock_session = AsyncMock()
        mock_session.execute.side_effect = [mock_no_version_result, mock_path_result]

        mock_cm = AsyncMock()
        mock_cm.__aenter__.return_value = mock_session
        mock_factory = MagicMock(return_value=mock_cm)
        monkeypatch.setattr("api.database.async_session_factory", mock_factory)

        result = await check_dedup("ws1", "hash456", "existing_path.pdf")
        assert result == "path-match-uuid"

    async def test_check_dedup_no_match(self, monkeypatch):
        from api.ingestion.dedup import check_dedup

        mock_no_result = MagicMock()
        mock_no_result.scalar_one_or_none.return_value = None

        mock_session = AsyncMock()
        mock_session.execute.return_value = mock_no_result

        mock_cm = AsyncMock()
        mock_cm.__aenter__.return_value = mock_session
        mock_factory = MagicMock(return_value=mock_cm)
        monkeypatch.setattr("api.database.async_session_factory", mock_factory)

        result = await check_dedup("ws1", "hash789", "new_file.pdf")
        assert result is None

    async def test_check_dedup_import_error(self, monkeypatch):
        from api.ingestion.dedup import check_dedup

        monkeypatch.setitem(sys.modules, "api.database", None)

        result = await check_dedup("ws1", "hash", "my_duplicate_file.pdf")
        assert result == "existing_doc_id_123"

        result = await check_dedup("ws1", "hash", "unique_file.pdf")
        assert result is None

    async def test_check_dedup_query_error(self, monkeypatch):
        from api.ingestion.dedup import check_dedup

        mock_session = AsyncMock()
        mock_session.execute.side_effect = Exception("DB connection lost")

        mock_cm = AsyncMock()
        mock_cm.__aenter__.return_value = mock_session
        mock_factory = MagicMock(return_value=mock_cm)
        monkeypatch.setattr("api.database.async_session_factory", mock_factory)

        result = await check_dedup("ws1", "hash", "my_duplicate_file.pdf")
        assert result == "existing_doc_id_123"

        result = await check_dedup("ws1", "hash", "unique_file.pdf")
        assert result is None

    async def test_fallback_dedup_duplicate(self):
        from api.ingestion.dedup import _fallback_dedup

        result = _fallback_dedup("ws1", "hash", "my_duplicate_file.pdf")
        assert result == "existing_doc_id_123"

    async def test_fallback_dedup_no_duplicate(self):
        from api.ingestion.dedup import _fallback_dedup

        result = _fallback_dedup("ws1", "hash", "unique_file.pdf")
        assert result is None


class TestPipeline:

    async def test_pipeline_new_doc(self, monkeypatch):
        from unittest.mock import AsyncMock, MagicMock
        from api.ingestion.parsers import ParsedDocument
        from api.ingestion.pipeline import run_pipeline

        async def mock_parse(filename, content):
            return ParsedDocument("New content", {"format": "markdown", "word_count": 2})

        async def mock_dedup(ws_id, content_hash, filename):
            return None

        mock_session = AsyncMock()
        mock_session.__aenter__ = AsyncMock(return_value=mock_session)
        mock_session.__aexit__ = AsyncMock(return_value=False)
        mock_session.begin = MagicMock(return_value=mock_session)
        mock_factory = MagicMock(return_value=mock_session)
        monkeypatch.setattr("api.ingestion.pipeline.parse_document", mock_parse)
        monkeypatch.setattr("api.ingestion.pipeline.check_dedup", mock_dedup)
        monkeypatch.setattr("api.ingestion.pipeline.async_session_factory", mock_factory)
        monkeypatch.setattr("api.ingestion.pipeline.chunk_text", lambda **kwargs: [])

        result = await run_pipeline("00000000-0000-0000-0000-000000000001", "new.md", b"# New")
        assert result["status"] == "success"
        assert "document_id" in result
        assert "version_id" in result
        assert result["metadata"]["format"] == "markdown"

    async def test_pipeline_existing_doc(self, monkeypatch):
        from unittest.mock import AsyncMock, MagicMock, PropertyMock
        from api.ingestion.parsers import ParsedDocument
        from api.ingestion.pipeline import run_pipeline

        async def mock_parse(filename, content):
            return ParsedDocument("Existing content", {"format": "pdf", "pages": 1, "word_count": 2})

        async def mock_dedup(ws_id, content_hash, filename):
            return "00000000-0000-0000-0000-000000000789"

        mock_doc = MagicMock()
        mock_doc.id = "00000000-0000-0000-0000-000000000789"
        mock_doc.metadata_ = {}
        mock_doc.updated_at = None

        mock_scalar_result = MagicMock()
        mock_scalar_result.scalar_one_or_none.return_value = mock_doc
        mock_scalar_result.scalar.return_value = 1

        mock_result = MagicMock()
        mock_result.scalar_one_or_none.return_value = mock_doc

        mock_version_result = MagicMock()
        mock_version_result.scalar.return_value = 1

        call_count = 0
        async def mock_execute(query):
            nonlocal call_count
            call_count += 1
            if call_count == 1:
                return mock_result
            return mock_version_result

        mock_session = AsyncMock()
        mock_session.__aenter__ = AsyncMock(return_value=mock_session)
        mock_session.__aexit__ = AsyncMock(return_value=False)
        mock_session.begin = MagicMock(return_value=mock_session)
        mock_session.execute = mock_execute
        mock_factory = MagicMock(return_value=mock_session)
        monkeypatch.setattr("api.ingestion.pipeline.parse_document", mock_parse)
        monkeypatch.setattr("api.ingestion.pipeline.check_dedup", mock_dedup)
        monkeypatch.setattr("api.ingestion.pipeline.async_session_factory", mock_factory)
        monkeypatch.setattr("api.ingestion.pipeline.chunk_text", lambda **kwargs: [])

        result = await run_pipeline("00000000-0000-0000-0000-000000000001", "existing.pdf", b"data")
        assert result["status"] == "success"
        assert "document_id" in result
        assert "version_id" in result

    async def test_pipeline_unsupported_format(self, monkeypatch):
        from api.ingestion.parsers import UnsupportedFormatError
        from api.ingestion.pipeline import run_pipeline

        async def mock_parse_error(filename, content):
            raise UnsupportedFormatError("No parser for .xyz")

        monkeypatch.setattr("api.ingestion.pipeline.parse_document", mock_parse_error)

        result = await run_pipeline("ws1", "bad.xyz", b"data")
        assert result["status"] == "error"
        assert "No parser for" in result["reason"]

    async def test_pipeline_generic_exception(self, monkeypatch):
        from api.ingestion.pipeline import run_pipeline

        async def mock_parse_error(filename, content):
            raise ValueError("Unexpected parse failure")

        monkeypatch.setattr("api.ingestion.pipeline.parse_document", mock_parse_error)

        result = await run_pipeline("ws1", "broken.pdf", b"data")
        assert result["status"] == "error"
        assert "Unexpected parse failure" in result["reason"]
