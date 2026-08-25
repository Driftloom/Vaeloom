"""
Document Builder Engine — converts structured resume JSON into styled artifacts:
  PDF  (Playwright Chromium print-to-PDF, pixel-perfect HTML/CSS)
  DOCX (python-docx, ATS-parseable editable Word)
  HTML (live preview for the web UI)

Also builds: matching cover letters, interview cheat-sheets, portfolio markdown.

Playwright is a lazy singleton — chromium launches on first compile and is reused.
If browsers are not installed the caller gets PlaywrightUnavailableError so the
API can return 503 with a setup hint instead of crashing.
"""
import asyncio
import contextlib
import io
import logging
import re
from dataclasses import dataclass

from api.services.resume_templates import normalize_resume_content, resume_templates

logger = logging.getLogger(__name__)

MAX_FIT_ATTEMPTS = 3
FIT_SCALE_STEP = 0.9  # shrink base font 10% per overflow pass


class DocumentBuilderError(Exception):
    """Rendering failed for a content/structural reason."""


class PlaywrightUnavailableError(DocumentBuilderError):
    """Chromium is not installed — operator must run `playwright install chromium`."""


@dataclass
class CompiledDocument:
    data: bytes
    media_type: str
    extension: str


class _PlaywrightManager:
    def __init__(self) -> None:
        self._lock = asyncio.Lock()
        self._playwright = None
        self._browser = None
        self._unavailable = False

    async def render_pdf(self, html: str) -> bytes:
        if self._unavailable:
            raise PlaywrightUnavailableError(
                "Playwright Chromium unavailable — run `uv run playwright install chromium`"
            )
        async with self._lock:
            try:
                if self._browser is None:
                    from playwright.async_api import async_playwright

                    self._playwright = await async_playwright().start()
                    self._browser = await self._playwright.chromium.launch(headless=True)
                page = await self._browser.new_page()
                try:
                    await page.set_content(html, wait_until="load")
                    return await page.pdf(format="A4", print_background=True)
                finally:
                    await page.close()
            except Exception as e:  # noqa: BLE001 - convert low-level errors
                msg = str(e)
                if "Executable" in msg or "install" in msg.lower() and "chromium" in msg.lower():
                    self._unavailable = True
                    logger.warning("Playwright chromium missing; PDF disabled until installed")
                    raise PlaywrightUnavailableError(
                        "Playwright Chromium unavailable — run `uv run playwright install chromium`"
                    ) from e
                raise DocumentBuilderError(f"PDF rendering failed: {e}") from e

    async def count_pdf_pages(self, pdf_bytes: bytes) -> int:
        try:
            import pymupdf as fitz
        except ImportError:  # older PyMuPDF versions
            import fitz

        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            return doc.page_count

    async def fetch_page_text(self, url: str, timeout_ms: int = 20000) -> tuple[str, str]:
        """Navigate to `url` in headless Chromium and return (body_text, title).

        Shared by document compilation (Phase 1) and browser scraping tools
        (Phase 2). Raises PlaywrightUnavailableError when chromium is missing.
        """
        return await self.fetch_page_text_guarded(url, None, timeout_ms)

    async def fetch_page_text_guarded(self, url: str, route_guard=None,
                                      timeout_ms: int = 20000) -> tuple[str, str]:
        """Like fetch_page_text but intercepts every network request through
        `route_guard` (async handler: route -> continue_/abort) so redirects
        and subresources are re-validated against the SSRF policy."""
        if self._unavailable:
            raise PlaywrightUnavailableError(
                "Playwright Chromium unavailable — run `uv run playwright install chromium`"
            )
        async with self._lock:
            try:
                if self._browser is None:
                    from playwright.async_api import async_playwright

                    self._playwright = await async_playwright().start()
                    self._browser = await self._playwright.chromium.launch(headless=True)
                context = await self._browser.new_context()
                try:
                    if route_guard is not None:
                        await context.route("**/*", route_guard)
                    page = await context.new_page()
                    try:
                        await page.goto(url, wait_until="domcontentloaded", timeout=timeout_ms)
                        title = await page.title()
                        text = await page.inner_text("body")
                        return (text or "").strip(), (title or "").strip()
                    finally:
                        with contextlib.suppress(Exception):
                            await page.close()
                finally:
                    with contextlib.suppress(Exception):
                        await context.close()
            except PlaywrightUnavailableError:
                raise
            except Exception as e:  # noqa: BLE001 - convert low-level errors
                msg = str(e)
                if "Executable" in msg or "install" in msg.lower() and "chromium" in msg.lower():
                    self._unavailable = True
                    logger.warning("Playwright chromium missing; browsing disabled until installed")
                    raise PlaywrightUnavailableError(
                        "Playwright Chromium unavailable — run `uv run playwright install chromium`"
                    ) from e
                raise DocumentBuilderError(f"Page fetch failed: {e}") from e

    async def close(self) -> None:
        async with self._lock:
            if self._browser:
                with contextlib.suppress(Exception):
                    await self._browser.close()
                self._browser = None
            if self._playwright:
                with contextlib.suppress(Exception):
                    await self._playwright.stop()
                self._playwright = None


_pw_manager = _PlaywrightManager()
# Public alias for other services (browser scraping tools) sharing one browser.
playwright_manager = _pw_manager


class DocumentBuilder:
    # ── Resume compilation ────────────────────────────────────────────
    async def compile_resume(self, content: dict, template_slug: str,
                             fmt: str = "pdf", max_pages: int = 2) -> CompiledDocument:
        tpl = resume_templates.get_template(template_slug)
        if tpl is None:
            raise ValueError(f"Unknown resume template: {template_slug}")
        html = resume_templates.render_resume_html(template_slug, content)

        if fmt == "html":
            return CompiledDocument(html.encode("utf-8"), "text/html", "html")
        if fmt == "docx":
            docx_bytes = self._resume_to_docx(normalize_resume_content(content))
            return CompiledDocument(docx_bytes,
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    "docx")
        if fmt != "pdf":
            raise ValueError(f"Unsupported format: {fmt}")

        # Page-fit loop: shrink typography until within target page budget
        scale = 1.0
        pdf = await _pw_manager.render_pdf(self._with_scale(html, scale))
        pages = await _pw_manager.count_pdf_pages(pdf)
        while pages > max_pages and scale > 0.6:
            scale *= FIT_SCALE_STEP
            pdf = await _pw_manager.render_pdf(self._with_scale(html, scale))
            pages = await _pw_manager.count_pdf_pages(pdf)
        return CompiledDocument(pdf, "application/pdf", "pdf")

    @staticmethod
    def _with_scale(html: str, scale: float) -> str:
        font_size_pt = round(10.5 * scale, 2)
        style_tag = (
            f"<style>body {{ font-size: {font_size_pt}pt !important; "
            f"line-height: {round(1.45 * scale, 2)} !important; }}</style>"
        )
        if "</head>" in html:
            return html.replace("</head>", f"{style_tag}</head>", 1)
        return style_tag + html

    # ── Cover letter ──────────────────────────────────────────────────
    async def compile_cover_letter(self, content: dict, body: str, template_slug: str,
                                   recipient: str | None = None, company: str | None = None,
                                   role: str | None = None, fmt: str = "pdf") -> CompiledDocument:
        html = resume_templates.render_cover_letter_html(
            template_slug, content, body, recipient=recipient, company=company, role=role
        )
        if fmt == "html":
            return CompiledDocument(html.encode("utf-8"), "text/html", "html")
        if fmt == "docx":
            paras = [p.strip() for p in body.split("\n\n") if p.strip()]
            data = normalize_resume_content(content)
            return CompiledDocument(self._letter_to_docx(data, paras),
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    "docx")
        if fmt != "pdf":
            raise ValueError(f"Unsupported format: {fmt}")
        pdf = await _pw_manager.render_pdf(html)
        return CompiledDocument(pdf, "application/pdf", "pdf")

    # ── Interview cheat-sheet ────────────────────────────────────────
    async def compile_cheatsheet(self, content: dict) -> CompiledDocument:
        html = resume_templates.render_cheatsheet_html(content)
        pdf = await _pw_manager.render_pdf(html)
        return CompiledDocument(pdf, "application/pdf", "pdf")

    # ── Portfolio markdown export ────────────────────────────────────
    @staticmethod
    def export_portfolio_markdown(content: dict) -> str:
        d = normalize_resume_content(content)
        lines: list[str] = [f"# {d['name']}", ""]
        if d["title"]:
            lines += [f"**{d['title']}**", ""]
        if d["summary"]:
            lines += [d["summary"], ""]
        links = [f"[{k.title()}]({v})" for k, v in d["links"].items() if v]
        if links:
            lines += [" | ".join(links), ""]
        if d["projects"]:
            lines += ["## Projects", ""]
            for p in d["projects"]:
                lines.append(f"### {p['name']}" if p["name"] else "### Project")
                if p["link"]:
                    lines.append(f"[Link]({p['link']})")
                if p["description"]:
                    lines += ["", p["description"]]
                for h in p["highlights"]:
                    lines.append(f"- {h}")
                lines.append("")
        if d["skills"]:
            lines += ["## Skills", ""]
            for g in d["skills"]:
                lines.append(f"- **{g['category']}:** " + ", ".join(i["name"] for i in g["items"]))
            lines.append("")
        if d["experience"]:
            lines += ["## Experience", ""]
            for e in d["experience"]:
                period = f" ({e['start']}–{e['end'] or 'Present'})" if e["start"] else ""
                lines.append(f"### {e['role']} @ {e['company']}{period}")
                for b in e["bullets"]:
                    lines.append(f"- {b}")
                lines.append("")
        return "\n".join(lines).strip() + "\n"

    # ── DOCX builders (ATS-parseable: standard headings, no tables/textboxes) ──
    @staticmethod
    def _resume_to_docx(d: dict) -> bytes:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor

        doc = DocxDocument()
        accent = RGBColor(0x1F, 0x3A, 0x5F)

        name_p = doc.add_paragraph()
        run = name_p.add_run(d["name"].upper())
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = accent

        contact_bits = [d["email"], d["phone"], d["location"]]
        contact_bits += [v for v in d["links"].values() if v]
        if any(contact_bits):
            doc.add_paragraph(" | ".join(b for b in contact_bits if b)).runs[0].font.size = Pt(9.5)
        if d["title"]:
            t = doc.add_paragraph()
            tr = t.add_run(d["title"])
            tr.font.size = Pt(12)
            tr.font.color.rgb = accent

        def heading(text: str) -> None:
            h = doc.add_heading(text, level=2)
            for r in h.runs:
                r.font.color.rgb = accent
                r.font.size = Pt(12)

        def bullets(items: list[str], style: str = "List Bullet") -> None:
            for item in items:
                p = doc.add_paragraph(item, style=style)
                p.paragraph_format.space_after = Pt(2)

        if d["summary"]:
            heading("Professional Summary")
            doc.add_paragraph(d["summary"])

        if d["experience"]:
            heading("Professional Experience")
            for e in d["experience"]:
                head = doc.add_paragraph()
                r = head.add_run(e["role"])
                r.bold = True
                meta = f"{e['company']}"
                if e["location"]:
                    meta += f", {e['location']}"
                if e["start"]:
                    meta += f"  |  {e['start']} – {e['end'] or 'Present'}"
                m = head.add_run("\n" + meta)
                m.italic = True
                m.font.size = Pt(9.5)
                if e["bullets"]:
                    bullets(e["bullets"])

        if d["education"]:
            heading("Education")
            for ed in d["education"]:
                p = doc.add_paragraph()
                r = p.add_run(ed["degree"])
                r.bold = True
                if ed["institution"]:
                    p.add_run(f"\n{ed['institution']}")
                if ed["details"]:
                    p.add_run(f" — {ed['details']}")

        if d["skills"]:
            heading("Skills")
            for g in d["skills"]:
                p = doc.add_paragraph()
                r = p.add_run(f"{g['category']}: ")
                r.bold = True
                p.add_run(", ".join(i["name"] for i in g["items"]))

        if d["projects"]:
            heading("Projects")
            for pr in d["projects"]:
                p = doc.add_paragraph()
                p.add_run(pr["name"]).bold = True
                if pr["description"]:
                    doc.add_paragraph(pr["description"])
                if pr["highlights"]:
                    bullets(pr["highlights"])

        if d["certifications"]:
            heading("Certifications")
            bullets(d["certifications"], style="List Bullet")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def _letter_to_docx(d: dict, paragraphs: list[str]) -> bytes:
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph(d["name"].upper()).runs[0].bold = True
        contact_bits = [b for b in (d["email"], d["phone"]) if b]
        if contact_bits:
            doc.add_paragraph(" | ".join(contact_bits))
        doc.add_paragraph("")
        for p in paragraphs:
            doc.add_paragraph(p)
        doc.add_paragraph("")
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph(d["name"])
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


document_builder = DocumentBuilder()

_SLUG_RE = re.compile(r"[^a-zA-Z0-9-_]")


def safe_filename(text: str, fallback: str = "document") -> str:
    cleaned = _SLUG_RE.sub("-", (text or "").strip())[:80]
    cleaned = re.sub(r"-{2,}", "-", cleaned).strip("-")
    return cleaned or fallback
